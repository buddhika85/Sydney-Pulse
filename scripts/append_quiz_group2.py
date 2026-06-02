"""Append SP1-14 Group 2 (TfNSW Client) quiz section to the Word doc.
One-shot script; safe to delete after use."""
from docx import Document

DOC_PATH = r"C:\BUDDHIKA\SydPulse-P6\SP1-14-Quiz-VehicleUpdate-ServiceAlert.docx"

SECTION_TITLE = "SP1-14: TfNSW Client (TfNswOptions, ITfNswFeedClient, TfNswFeedClient)"

QA = [
    {
        "q": (
            "TfNswFeedClient is registered as a singleton in Program.cs. Why is singleton the "
            "right lifetime here, what specifically would break (or get worse) if you changed "
            "it to transient, and what would change if you made it scoped?"
        ),
        "a": [
            "Singleton lifetime in DI is per Function App process, not per Function or per invocation. Within one process, every invocation across every function (PollerFunction, RoutesFunction, HTTP requests) shares the same TfNswFeedClient instance.",
            "The class holds mutable cross-invocation state: the per-mode route cache (_routeCache / _cacheExpiry) and the SemaphoreSlim (_cacheLock). Both are only useful if shared across invocations.",
            "Transient would create a fresh instance per injection: each invocation starts with an empty cache, re-downloads the 20-50 MB GTFS-static ZIP every 30 seconds, and the per-instance semaphore coordinates nothing.",
            "Scoped in Azure Functions isolated worker means per-function-invocation, which is effectively identical to transient for this class. Scoped exists as a concept (think DbContext per HTTP request) but buys nothing here.",
            "If Azure scales out to multiple Function App instances (separate processes) under load, each process gets its own singleton, its own cache, and its own lock. SemaphoreSlim cannot coordinate across processes - it is in-process only.",
        ],
    },
    {
        "q": (
            "Three concurrent GetRoutesAsync(\"buses\") calls land at the exact instant the "
            "bus cache has just expired. Walk through what each call does line-by-line, and "
            "explain how many GTFS-static ZIP downloads end up happening. Pay attention to why "
            "the expiry check is performed twice."
        ),
        "a": [
            "All three threads pass the fast-path check on line 126 (cache expired) without acquiring the lock.",
            "All three call await _cacheLock.WaitAsync(...). The semaphore has capacity 1, so thread A acquires the lock immediately; threads B and C async-wait.",
            "Thread A enters the try block, re-checks expiry on line 134 (still expired), calls FetchAndParseStaticRoutesAsync, downloads the ZIP, populates _routeCache[mode] and _cacheExpiry[mode], returns the dictionary. The finally block calls Release().",
            "Thread B is woken next, re-checks expiry on line 134 - now non-expired because A refreshed it - returns the cached dictionary. Finally block calls Release().",
            "Thread C does the same as B.",
            "Net result: exactly one GTFS-static ZIP download across the three threads. This is the double-checked locking pattern: fast-path check outside the lock for warm-cache reads, authoritative re-check inside the lock to absorb threads that queued during a refresh.",
            "C# semantics: return inside try still executes finally, so all three threads release the semaphore correctly.",
            "Subtle issue 1 - dictionaries are not thread-safe. Line 126 reads _cacheExpiry outside the lock; line 139 writes to it inside the lock. The race window is microseconds, once per hour per mode, so observable corruption is near zero, but ConcurrentDictionary would be the textbook fix.",
            "Subtle issue 2 - one lock covers all modes. If 'buses' and 'ferries' expire simultaneously, they refresh serially even though independent. Per-mode locks would allow parallel refreshes. Acceptable trade-off at our scale.",
        ],
    },
    {
        "q": (
            "TfNSW has a 5 rps rate limit. PollerFunction iterates six modes back-to-back and "
            "the 5th call returns HTTP 429. The retry policy lives on the named TfNsw HttpClient "
            "in Program.cs, not in GetVehiclePositionsAsync. Who sees the 429 first? Does the "
            "FeedClient know about retries? Why is this layering preferable to wrapping retries "
            "inside the FeedClient itself? How is a circuit breaker different from a retry?"
        ),
        "a": [
            "Polly slots into the HttpClient pipeline as a DelegatingHandler (or modern resilience handler). The handler inspects the 429 response, decides it is transient, waits per the backoff policy, and re-sends through the pipeline.",
            "GetVehiclePositionsAsync does NOT see retries happening on the happy path. It awaits SendAsync once and receives the final response (success or terminal failure).",
            "It DOES see terminal failure: if all retries are exhausted, response.EnsureSuccessStatusCode() in FetchBytesAsync throws HttpRequestException, which propagates to the FeedClient caller.",
            "Why HttpClient-layer is preferable - single responsibility: FeedClient owns protobuf decoding and cache management. HTTP transport reliability is a separate concern.",
            "Why HttpClient-layer is preferable - pipeline composability: the standard resilience pipeline isn't just retries. It bundles retry + timeout + circuit breaker + rate limiter. These wrap the transport call and can't be bolted onto a business method.",
            "Why HttpClient-layer is preferable - the double-retry hazard: if retries lived inside GetVehiclePositionsAsync as well, they would multiply with Polly's retries (e.g., 3 outer x 3 Polly = up to 9 attempts). This would blow past TfNSW's 5 rps limit and might trip the breaker on a self-induced failure flood. Retries belong at exactly one layer.",
            "Retry semantics: per-request, each call gets its own retry budget. Use for transient single-request failures (network blip, single 429, single 5xx).",
            "Circuit breaker semantics: per-resource, tracked across many requests. Counts failure rate (e.g., 10 failures in 30s). When threshold trips, breaker OPENS - subsequent calls fail instantly without hitting the network (fail-fast). After a cool-down, breaker enters HALF-OPEN - one trial call; success closes it, failure re-opens.",
            "They work together: retry handles per-request transients; breaker handles 'this downstream is on fire, stop dialling for a bit'.",
        ],
    },
    {
        "q": (
            "When GetVehiclePositionsAsync calls /v2/gtfs/vehiclepos/{mode}, what format does "
            "the response arrive in over the wire, and how is it decoded? Why does the method "
            "also call GetRoutesAsync(mode)? CLAUDE.md flags that route_id (e.g. NTH_1a) is "
            "internal while route_short_name (e.g. T1) is user-facing - where does that "
            "distinction appear in this method, and what would break in the system if route_id "
            "leaked instead of being resolved to route_short_name?"
        ),
        "a": [
            "Response body is binary protobuf encoding (GTFS-Realtime spec), not JSON. Protobuf is more compact and faster to decode than JSON.",
            "Decoding: FeedMessage.Parser.ParseFrom(bytes) parses raw bytes into a strongly-typed FeedMessage object.",
            "Schema source-of-truth: GtfsRealtime.proto in SydneyPulse.Core/TfNsw/. The Grpc.Tools MSBuild target generates C# classes at BUILD time (not runtime) - generated .cs files land in obj/ and are compiled into the assembly like any other source. By the time code runs, FeedMessage/VehiclePosition/Alert are ordinary compiled CLR types.",
            "Why also call GetRoutesAsync: GTFS is two separate specs. GTFS-Static is the catalogue (routes.txt, stops.txt, schedules) - slow-changing, distributed as a ZIP of CSVs (~20-50 MB compressed, ~150 MB uncompressed). GTFS-Realtime is the current state (vehicle positions, alerts, trip updates) - small, fast-changing, protobuf streams.",
            "The realtime feed only contains route_id. To populate route_short_name, route_long_name, and route_color on VehicleUpdate, the client must merge in static catalogue data. Caching the static feed for 1 hour avoids re-downloading 20-50 MB every 30 seconds.",
            "Distinction in code (line 75): RouteShortName: route?.ShortName ?? routeId - try static lookup, fall back to route_id if missing (graceful degradation when static cache lags realtime).",
            "What breaks if route_id leaks - the Cosmos partition key is /routeShortName (ADR-0002). Vehicles on the T1 line carrying different route_ids (e.g., NTH_1a, STH_2b) would shard to different partitions; identical-route vehicles no longer co-locate.",
            "What breaks if route_id leaks - GET /api/vehicles?routeShortName=T1 (a partition-key query in VehiclesFunction) would return nothing, because the partition would contain raw route_ids, not T1.",
            "What breaks if route_id leaks - cross-mode aggregation and alert correlation break. The whole system is keyed on route_short_name.",
            "What breaks if route_id leaks - frontend would show ugly internal IDs to humans. Readability gone.",
        ],
    },
    {
        "q": (
            "TfNswFeedClient ships with an ITfNswFeedClient interface, and PollerFunction + "
            "RoutesFunction depend on the interface, not the concrete class. What practical "
            "benefit does this give you? If you deleted the interface and injected the concrete "
            "class directly, what specifically becomes harder? Is the interface always a "
            "defensible choice, or does the modern C# pendulum push back on it?"
        ),
        "a": [
            "The interface decouples consumers (PollerFunction, RoutesFunction) from the concrete implementation, enabling Dependency Injection and Inversion of Control.",
            "Mechanical benefit - Moq can only intercept calls on interface members or virtual methods. Without ITfNswFeedClient, unit tests in PollerFunctionTests / RoutesFunctionTests would either need every public method on TfNswFeedClient to be virtual (intrusive change just to enable testing) or hand-rolled test doubles.",
            "Architectural benefit - tests against the concrete client would actually hit the TfNSW HTTP API. Those would no longer be unit tests; they would be integration tests - slow, brittle, network-dependent, rate-limited.",
            "Polymorphism - enables swapping implementations without touching consumers (e.g., a chaos-testing fake feed client, or a v2 TfNswFeedClient deployed alongside v1).",
            "Modern C# critique: interfaces are NOT always free. The 2020s pendulum has swung against 'one interface per class by reflex' because interfaces have maintenance cost - drift between interface and impl, duplicate signatures, IDE noise.",
            "Extract an interface when - the class crosses a process boundary (HTTP, DB, message bus) and needs stubbing for unit tests; multiple implementations exist or are likely; side effects need to be controlled in tests.",
            "Skip the interface when - pure data / DTOs / records (e.g., VehicleUpdate); constants / static utilities (e.g., FunctionConstants); one concrete implementation that you don't plan to swap, with no I/O.",
            "TfNswFeedClient is a textbook case for an interface: it hits a remote HTTP API and every consumer's unit test needs to stub the network out. The maintenance cost is justified by the testability win.",
        ],
    },
    {
        "q": (
            "FetchBytesAsync reads _options.ApiKey and adds it to the Authorization header. "
            "Trace that secret backwards along two paths: where does it physically live in "
            "local dev, and how does it reach _options.ApiKey? Where does it live in Azure, "
            "and how does it reach _options.ApiKey? Who is authorised to read it in Azure, "
            "and how is that authorisation proven without storing a password? Finally, what "
            "threat does the Key-Vault-via-Managed-Identity design defend against that a plain "
            "app setting containing the key would not?"
        ),
        "a": [
            "Local dev path - the secret physically lives in functions/SydneyPulse.Functions/local.settings.json (gitignored, and denied to Claude tools via .claude/settings.json).",
            "Local dev path - file format in Functions isolated worker is a flat Values envelope with double-underscore for nesting: 'TfNsw__ApiKey': '...'. This is different from ASP.NET Core's appsettings.json nested-JSON style.",
            "Local dev path - Functions host reads local.settings.json on startup and exposes each Values entry as a process environment variable.",
            "Local dev path - Program.cs calls services.Configure<TfNswOptions>(context.Configuration.GetSection(TfNswOptions.SectionName)). The configuration system reads the env var via the ':' <-> '__' conversion convention and binds it into a strongly-typed TfNswOptions.",
            "Local dev path - TfNswFeedClient constructor receives IOptions<TfNswOptions> via DI; _options.ApiKey now holds the secret value; FetchBytesAsync places it in the Authorization header.",
            "Azure path - the secret lives in Azure Key Vault as the TfNswApiKey secret. Encrypted at rest and in transit.",
            "Azure path - the Function App's app setting for TfNsw__ApiKey holds a Key Vault reference: @Microsoft.KeyVault(SecretUri=https://<vault>.vault.azure.net/secrets/TfNswApiKey) or the VaultName/SecretName form.",
            "Azure path - the Function App's system-assigned Managed Identity is granted the 'Key Vault Secrets User' role (data-plane role) on the vault, configured in role-assignments.bicep.",
            "Azure startup resolution hop-by-hop: (1) Functions host iterates app settings, sees value starts with @Microsoft.KeyVault(...). (2) Host requests an access token from IMDS (Instance Metadata Service at 169.254.169.254) using the MI. (3) Entra ID issues a token scoped to https://vault.azure.net. (4) Host calls Key Vault REST with the token, fetches the secret. (5) Host substitutes the resolved value as the runtime env var. (6) IConfiguration -> IOptions<TfNswOptions> -> _options.ApiKey - same binding as local. All complete before any user code runs.",
            "Rotation: update the secret in Key Vault; Function App picks up the new value on next restart (or within ~24h periodic refresh when using the unversioned SecretUri= form). One update refreshes all consumers - no code change, no redeploy.",
            "Threat defended #1 - secrets never appear in source code or git. Eliminates the largest historical secret-leak vector.",
            "Threat defended #2 - separate access boundary. Plain Function App settings ARE encrypted at rest, but visible in plaintext to anyone with Contributor RBAC on the Function App (many engineers usually have this). Key Vault enforces a SEPARATE RBAC boundary - only identities with explicit Key Vault Secrets User on the vault can read, even if they hold Function App Contributor. This is least privilege in action.",
            "Threat defended #3 - audit trail. Every Key Vault secret read is logged to Azure Monitor (who, when). Plain app-setting reads are not audited.",
            "Threat defended #4 - rotation discipline. One Key Vault update propagates to all consumers automatically; plain app settings require finding every copy and restarting each consumer.",
            "Threat defended #5 - no shared credential to steal. Managed Identity is credential-less; Azure exchanges short-lived Entra tokens through IMDS. There is no password, connection string, or SAS to leak.",
        ],
    },
]


def main() -> None:
    doc = Document(DOC_PATH)

    # Drop trailing blank Normal paragraphs before appending so the new section
    # starts cleanly after the previous one.
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)

    doc.add_paragraph("")  # one-line separator
    doc.add_paragraph(SECTION_TITLE, style="Heading 1")

    for i, item in enumerate(QA, start=1):
        doc.add_paragraph(f"Question {i}", style="Heading 2")
        doc.add_paragraph(item["q"], style="Normal")
        doc.add_paragraph("Model Answer", style="Heading 3")
        for bullet in item["a"]:
            doc.add_paragraph(bullet, style="List Bullet")
        doc.add_paragraph("")  # separator after each Q

    doc.save(DOC_PATH)
    print(f"Saved {DOC_PATH}")
    print(f"Added section: {SECTION_TITLE} with {len(QA)} questions.")


if __name__ == "__main__":
    main()
