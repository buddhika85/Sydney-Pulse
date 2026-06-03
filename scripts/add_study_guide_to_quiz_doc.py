# Inserts a "How to Study This Document" guide at the top of the SP1-14 quiz Word doc.
# Already executed on 2026-06-03; retained as audit trail.
# Re-running will duplicate the guide section; remove it in Word first if you need to re-run.

from docx import Document
from pathlib import Path

DOC_PATH = Path(r"C:\BUDDHIKA\SydPulse-P6\SP1-14-Quiz-VehicleUpdate-ServiceAlert.docx")

# Each item: (style_name, text). Inserted at the very top of the doc in this order.
# "List Number" gives an auto-numbered list; "List Bullet" gives bullet points.
# If a style is missing from the doc, python-docx falls back to Word's built-in.
GUIDE_CONTENT = [
    ("Heading 1", "How to Study This Document"),
    ("Normal", "This document is a reference, not a textbook. Don't read it cover-to-cover."),
    ("Normal", "The single rule: test yourself before reading the answer. Re-reading without active recall feels productive but barely works."),

    ("Heading 2", "Daily - 10 minutes"),
    ("List Number", "Open the doc at a random question."),
    ("List Number", "Cover the Model Answer with your hand or scroll past it."),
    ("List Number", "Read just the question text."),
    ("List Number", "Speak your answer out loud as if you're in an interview."),
    ("List Number", "Uncover the model answer. Mentally note one thing you missed."),
    ("List Number", "Stop. Don't re-read. Move on."),

    ("Heading 2", "Weekly - 30 minutes"),
    ("Normal", "Walk through one whole group (e.g., all PollerFunction questions). Time yourself: aim for under 90 seconds per question. This builds interview pace."),

    ("Heading 2", "Pre-interview - night before only, 15 minutes"),
    ("Normal", "Scan only the sections relevant to the company's tech stack. Azure-heavy shop -> review Q5 and TfNSW client. Event-driven systems -> review Q6. Don't re-read everything. You're priming, not learning."),

    ("Heading 2", "What NOT to do"),
    ("List Bullet", "Don't try to memorize word-for-word. Interviewers spot rehearsed answers."),
    ("List Bullet", "Don't read it like a textbook front to back."),
    ("List Bullet", "Don't make separate summary notes from this doc. The doc IS the summary."),
    ("List Bullet", "Don't re-read on the day of the interview. Spikes anxiety. Sleep does the consolidation."),

    ("Heading 2", "The minimum that beats most candidates"),
    ("Normal", "The night before any technical interview: 5 random questions, speak your answer out loud, peek at the model. 15 minutes total. That alone beats 90% of candidates who just re-read their resume."),

    # Visual spacer between the guide and the first existing quiz section.
    ("Normal", ""),
    ("Normal", ""),
]


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"Doc not found: {DOC_PATH}")

    doc = Document(str(DOC_PATH))

    if not doc.paragraphs:
        raise SystemExit("Doc has no paragraphs - cannot determine insertion point.")

    # Capture the XML element of the original first paragraph BEFORE adding anything.
    # All new content will be moved to immediately before this element.
    original_first_elem = doc.paragraphs[0]._element

    # python-docx has no native "insert at front" API. Standard workaround:
    # create the new paragraphs at the end of the doc, then use lxml addprevious()
    # to move each one to immediately before the original first paragraph.
    new_elements = []
    for style_name, text in GUIDE_CONTENT:
        try:
            p = doc.add_paragraph(text, style=style_name)
        except KeyError:
            # Style missing - fall back to Normal so the script doesn't fail.
            print(f"WARNING: style '{style_name}' not found; using Normal for: {text[:60]}")
            p = doc.add_paragraph(text, style="Normal")
        new_elements.append(p._element)

    # Move each new paragraph to immediately before the original first paragraph,
    # preserving insertion order (addprevious inserts directly before the target).
    for elem in new_elements:
        original_first_elem.addprevious(elem)

    doc.save(str(DOC_PATH))
    print(f"Inserted {len(GUIDE_CONTENT)} guide paragraphs at the top of {DOC_PATH.name}")


if __name__ == "__main__":
    main()
