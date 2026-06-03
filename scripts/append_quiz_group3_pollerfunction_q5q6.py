# Appends Q5, Q5 follow-up, and Q6 to the SP1-14 quiz Word doc.
# Companion to append_quiz_group3_pollerfunction.py (which appended Q1-Q4).
# Already executed on 2026-06-03; retained as audit trail.
# Re-running will duplicate content; delete the appended sections in Word first.

from docx import Document
from pathlib import Path

DOC_PATH = Path(r"C:\BUDDHIKA\SydPulse-P6\SP1-14-Quiz-VehicleUpdate-ServiceAlert.docx")

# Section title update: rename the PollerFunction Heading 1 to drop the "pending" suffix.
TITLE_OLD = "SP1-14: PollerFunction (Q1-Q4 of 6 - Q5-Q6 pending next session)"
TITLE_NEW = "SP1-14: PollerFunction (all 6 questions + 1 follow-up)"

# Each question block: heading + question text + list of model-answer bullets.
QUESTIONS = [
    {
        "heading": "Question 5 (a) - DefaultAzureCredential: how does it resolve locally vs in Azure?",
        "text": (
            "EventGridOptions has only TopicEndpoint - no API key, no SAS, no shared key. "
            "EventGridPublisherClient is registered with DefaultAzureCredential. How does this "
            "work locally vs in Azure (which credential is actually used in each environment)?"
        ),
        "answers": [
            "DefaultAzureCredential is not a binary local/Azure switch but an ordered chain of credential sources tried in sequence.",
            "Chain order: EnvironmentCredential -> WorkloadIdentityCredential -> ManagedIdentityCredential -> AzureCliCredential -> AzurePowerShellCredential -> AzureDeveloperCliCredential -> InteractiveBrowserCredential (disabled by default).",
            "In Azure (Function App): ManagedIdentityCredential succeeds first because IMDS responds at 169.254.169.254 in milliseconds. The Function App's system-assigned MI is the identity used.",
            "On a developer laptop: ManagedIdentityCredential step fails (IMDS not reachable, ~5s timeout), then AzureCliCredential succeeds using the developer's `az login` token.",
            "Same one-line registration works in both environments because the chain transparently selects whichever credential is available at the host.",
            "No secrets are needed in local.settings.json for the Event Grid TopicEndpoint - the endpoint URL itself is not a secret; auth is credential-based.",
            "Pro tip: the IMDS step on a laptop is slow to time out. Set AZURE_TOKEN_CREDENTIALS=AzureCliCredential or use DefaultAzureCredentialOptions.ExcludeManagedIdentityCredential = true to skip it during local dev.",
        ],
    },
    {
        "heading": "Question 5 (a) follow-up - Can we run this locally? Are we mocking Event Grid with Docker?",
        "text": (
            "Can we run this locally? When running, will it actually publish to the real Event Grid? "
            "Aren't we mocking Event Grid using Docker containers like the Azurite or Service Bus emulator?"
        ),
        "answers": [
            "Yes, you can run locally via `func start`. The host runs on localhost:7071 and the Timer trigger fires every 30s.",
            "The Poller hits the REAL sydney-pulse-events-dev Event Grid topic in Azure. No Microsoft-provided Event Grid emulator exists, and this project does not use a community simulator.",
            "What IS emulated locally: Azurite container for AzureWebJobsStorage (Functions runtime metadata only - timer schedules, dedupe state, etc.).",
            "What is NOT emulated for regular dev: Event Grid (no emulator exists), Cosmos DB (we point at real dev Cosmos), SignalR Service (real Free SKU per ADR-0008), Service Bus topic (real namespace per ADR-0003).",
            "Service Bus emulator container is wired up only for integration tests (per functions/CLAUDE.md), NOT for regular `func start` development loop.",
            "Practical implication - Cost: `func start` on your laptop = real dev-environment fan-out. Vehicles flow into real Cosmos, real Event Grid, real SignalR. At 30s cadence with 5 modes, leaving it running 24/7 costs real money per cost-model.md ($6-15/month baseline).",
            "Practical implication - Permissions: locally you authenticate as YOUR `az login` identity. That identity must have EventGrid Data Sender (or inherited Owner/Contributor on the dev subscription) on the dev topic, or SendEventsAsync returns 403.",
            "Why no Event Grid emulator: Microsoft's stance is that Event Grid is a routing layer (retries, dead-lettering, fan-out filters) - emulating it means emulating most of the product semantics. They've never shipped one.",
            "Stop the host when done with local testing to avoid burning real Cosmos RUs and Event Grid quota.",
        ],
    },
    {
        "heading": "Question 5 (b) - Role grant and Bicep pattern",
        "text": (
            "For the Function App in Azure to publish to the Event Grid topic, it needs an authorization "
            "grant on the topic. What role is granted, what principal is it granted to, where in the "
            "codebase is that grant wired up, and why use a role assignment rather than a connection "
            "string or SAS token?"
        ),
        "answers": [
            "Role: EventGrid Data Sender (built-in Azure role). Built-in role GUID: d5a91429-5739-47e2-a06b-3470a27159e7.",
            "Principal: the Function App's system-assigned managed identity (MI). MIs are a special kind of service principal automatically tied to a resource's lifecycle.",
            "Three identity flavours in Azure: App Registration (manual SP with client secret), User-assigned MI (standalone, attachable to multiple resources), System-assigned MI (auto-created per resource). This project uses system-assigned per SP1-03.",
            "If you delete and recreate the Function App, the system-assigned MI is NEW - you have to re-grant the role.",
            "Location in codebase: infra/modules/role-assignments.bicep (the dedicated RBAC module).",
            "Why role assignments live in their own Bicep module: role assignment `name` values must be GUIDs and pre-computable at compile time. Inlining in main.bicep triggers a Bicep compile error. Moving them to a scoped module avoids this.",
            "Bicep pattern shape: resource type 'Microsoft.Authorization/roleAssignments', scope set to target resource, name = guid(scope, principal, role) for deterministic redeploy, properties = roleDefinitionId (built via subscriptionResourceId), principalId, principalType: 'ServicePrincipal'.",
            "Why guid() for the name: role assignment names must be GUIDs (Azure requirement). guid() generates deterministic GUIDs so re-deploys don't create duplicates.",
            "Why role assignment beats connection string: connection strings grant full topic admin rights (not granular), must be stored securely (KV plumbing), must be rotated (lifecycle overhead), and expose leakage risk in source/logs.",
            "Why role assignment beats SAS: even though Event Grid does support SAS via aeg-sas-token header (developer initially thought SAS was storage-only - actually Event Grid and Service Bus both support SAS), SAS still requires the same secret-management overhead - issue, store, transmit, rotate.",
            "Beneficial properties of role assignment + MI: no secret to leak, auditable per-identity in Azure activity logs, zero rotation overhead (tokens minted on demand by AAD with ~1h TTL), identity tied to resource lifecycle.",
        ],
    },
    {
        "heading": "Question 5 (c) - Compare to TfNSW key story",
        "text": (
            "Compare the Event Grid auth story with the TfNSW API key story from Group 2 (key lives in "
            "Key Vault, retrieved via the Function App's MI). Same security pattern or different? What's "
            "the shared underlying primitive? Why can't we apply the 'no secret' pattern to TfNSW?"
        ),
        "answers": [
            "Authentication root: same in both - Function App's managed identity.",
            "Authorization mechanism: same - Azure RBAC. Event Grid: EventGrid Data Sender on the topic. Key Vault: Key Vault Secrets User on the vault.",
            "Use of MI differs: Event Grid is 1-hop (MI -> resource for data-plane access). TfNSW is 2-hop (MI -> Key Vault for secret retrieval -> static API key sent to TfNSW in-band).",
            "Shared underlying primitive: Managed Identity + Azure RBAC as a pair. MI gives who you are; RBAC defines what you can do.",
            "If you removed MI from the architecture - Event Grid path: would fall back to connection string / topic key / SAS in app settings or KV. The 'no secret' property is gone.",
            "If you removed MI from the architecture - TfNSW path: Function App couldn't authenticate to KV at all. Key would have to live in app settings as plain text or in source - every deployment script and log line becomes a potential leak vector.",
            "The fundamental asymmetry: Event Grid is within the Azure AD trust boundary; TfNSW is not. The 'no secret at all' pattern only works when the destination resource accepts Azure AD-issued tokens.",
            "Why TfNSW can't be credential-less: TfNSW is a third-party REST API with a single auth mechanism (apikey HTTP header carrying a static string). They are not a relying party for Azure AD. They don't speak OAuth in a way we can authenticate to.",
            "Mitigation when stuck with a secret (defense-in-depth): store in KV (not source, not app settings); gate access via KV RBAC (only Function App MI can read); audit via KV diagnostic logs; rotate periodically (KV versioning supports this).",
            "Token lifetime asymmetry: MI tokens are short-lived (~1h, auto-rotated). TfNSW API key is long-lived (manual rotation via TfNSW developer portal).",
            "Incident response asymmetry: leaked MI token expires in an hour anyway. Leaked TfNSW key requires manual rotation through TfNSW portal + KV update + accepted downtime window.",
            "Shorthand rule: you can only be 'credential-less' within an Azure AD trust boundary. Outside it, the best you can achieve is 'credential, but well-isolated'.",
        ],
    },
    {
        "heading": "Question 5 (d) - Failure modes",
        "text": (
            "What would break - and how would it manifest at runtime - if (i) you forgot to add the "
            "EventGrid Data Sender role assignment in Bicep, or (ii) you tried to run locally without "
            "`az login`? How would you distinguish the two failures?"
        ),
        "answers": [
            "Scenario (i) - missing role in Azure: Bicep deploy succeeds (resource shapes valid; only data-plane access is broken). Function App starts cleanly. First timer tick fires. SendEventsAsync calls Event Grid with the MI's AAD token. Event Grid checks RBAC, MI has no role, returns 403 Forbidden.",
            "Exception type for (i): Azure.RequestFailedException with Status = 403 and 'Forbidden' message.",
            "Functions runtime behaviour with Timer triggers: catches the exception, marks invocation as failed, does NOT retry (Timer triggers don't replay; next tick is 30s later). Result: one failed invocation every 30s until the role is added.",
            "App Insights surface for (i): requests table (success == false), exceptions table (full stack trace). KQL: requests | where name == 'Poller' and success == false | order by timestamp desc.",
            "Why Bicep deploy succeeded despite the bug: Bicep validates resource shapes, not data-plane access. Function App and topic both exist; nothing tells Bicep that the combination is broken until runtime.",
            "Side question (developer asked): why no try/catch in PollerFunction? Per functions/CLAUDE.md: 'Don't swallow exceptions. Functions runtime will retry the trigger based on the binding.' For Service Bus triggers, swallowing = message lost. For Event Grid triggers, swallowing = no retry. For Timer triggers, swallowing = invocation falsely reports success == true to App Insights -> broken alerting.",
            "When to use try/catch in a Function: partial-failure scenarios where you want to log + continue (e.g., one mode fails, still try others). Rule of thumb: log THEN rethrow, never log THEN swallow.",
            "Side question (developer asked): why not Serilog to text files? Functions Consumption hosts have no durable local filesystem. Hosts can cold-start, scale out, or relocate at any time; text files vanish.",
            "Even on Premium/Dedicated plans, multi-instance scale-out scatters logs across instances - no single file to read.",
            "App Insights is the right sink for Function logs: structured, queryable via KQL, configurable retention, sampling already wired (5% per host.json constraint).",
            "Serilog could still be used for its enrichers or structured features, but its sink would be Application Insights, not files. The built-in ILogger<T> already goes there via the Functions host's OpenTelemetry integration.",
            "Scenario (ii) - no `az login` locally: DefaultAzureCredential walks the chain. EnvironmentCredential fails (no env vars set). WorkloadIdentityCredential fails (not AKS). ManagedIdentityCredential calls IMDS at 169.254.169.254 - on a laptop this times out (~5s) before failing. This is why startup feels sluggish without `az login`.",
            "AzureCliCredential then runs `az account get-access-token` under the hood. Without `az login` (or with expired token), fails with message including 'Please run az login'.",
            "Subsequent steps (AzurePowerShellCredential, AzureDeveloperCliCredential) also fail.",
            "Thrown exception for (ii): AuthenticationFailedException (often wrapping multiple CredentialUnavailableException). Console output includes per-step failure summary.",
            "Functions host does NOT crash in scenario (ii): DI registers the EventGridPublisherClient lazily. The host starts fine; credential failure surfaces on the first publish attempt at the first timer tick.",
            "Pro tip: short-circuit the slow IMDS step in dev by setting AZURE_TOKEN_CREDENTIALS=AzureCliCredential or ExcludeManagedIdentityCredential = true.",
            "Distinguishing the two - the AuthN vs AuthZ frame: Scenario (i) authentication SUCCEEDED (got MI token), authorization FAILED (no role on topic). HTTP request to Event Grid WAS made; returned 403. Scenario (ii) authentication ITSELF failed (no credential available). HTTP request to Event Grid was NEVER made.",
            "One-line diagnostic: did an HTTP request to Event Grid actually go out? If yes (and got 403) -> role missing. If no (failed before HTTP layer) -> credential acquisition failed.",
            "Other tells: environment marker - (i) is Azure-only (locally your az login likely has Owner/Contributor inheritance on dev subscription, so RBAC passes by inheritance), (ii) is local-only (in Azure MI is always present).",
            "Stack trace depth tell: (i) starts in the SDK's HTTP response handler; (ii) starts at the credential layer, before any HTTP request is constructed.",
        ],
    },
    {
        "heading": "Question 6 (a) - The empty-batch guard",
        "text": (
            "PollerFunction has `if (vehicles.Count == 0) return;` in both publish methods. "
            "EventGridPublisherClient.SendEventsAsync(emptyList) could just be allowed to no-op. "
            "Why is the guard worth the line of code anyway?"
        ),
        "answers": [
            "CPU waste: invoking the publish path with nothing to publish wastes Function execution time on serialization and call setup.",
            "Network waste: empty SendEventsAsync would still make an HTTP call to Event Grid's REST endpoint.",
            "Direct cost: Event Grid charges per operation (~$0.60/million). Empty publishes still count. At 30s cadence x 5 modes x 2 event types = 86,400 calls/day. Half empty overnight = real money over a year.",
            "Observability hygiene: with the guard, every SendEventsAsync call in App Insights represents real data flowing. Without the guard, traces would have a noise floor of empty publishes polluting KQL queries like 'how many events per hour did we publish?'",
            "Latency: avoiding the unnecessary RPC reduces the function's overall execution time per tick.",
            "Behaviour pinned by test: RunAsync_WithEmptyFeeds_DoesNotCallSendEvents in PollerFunctionTests.cs. It's a deliberate contract, not just an optimization. A future refactor that removes the guard would fail the test and call it out.",
        ],
    },
    {
        "heading": "Question 6 (b) - When zero vehicles is legitimate",
        "text": (
            "The Poller iterates 5 transport modes (sydneytrains, ferries, buses, lightrail, nswtrains). "
            "When during a normal day would TfNSW legitimately return zero vehicles for one or more of "
            "these modes? List operational, scheduled, and environmental scenarios."
        ),
        "answers": [
            "Important distinction: 'feed responds with empty array' vs 'feed call fails' are different runtime behaviours. The latter throws HttpRequestException (and the Polly retry policy on the HttpClient handles transient failures). 'Endpoint down' / 'network issue' / 'API key rotation outage' fall in the failure bucket, not the legitimate empty bucket.",
            "Off-hours / overnight: ferries stop around midnight, light rail around 1am, most buses 1-4am. cost-model.md explicitly mentions stopping bus polling between 1-4am as a cost-saving lever.",
            "Severe weather: ferries suspended in high wind or storm; light rail suspended during lightning. Service paused for safety; feed responds normally but reports zero vehicles in service.",
            "Industrial action: train strikes mean no vehicles running. Feed publishes empty.",
            "Planned trackwork: weekend track closures on Sydneytrains / NSW trains lines = zero vehicles on the affected routes.",
            "Public holidays: reduced service or no service at all on some lines / modes.",
            "TfNSW backend maintenance windows: they sometimes push empty feeds rather than errors during their own maintenance windows - preferable from a consumer's perspective.",
            "Mode geographically tiny: light rail covers CBD/Inner West + Parramatta only. A single signal failure or incident takes all vehicles out briefly.",
            "Clarification on 'stopped vehicles' guess: stopped vehicles DO appear in the feed (with status STOPPED in GTFS-RT). 'Zero vehicles' means out of service, not 'not moving'.",
        ],
    },
    {
        "heading": "Question 6 (c) - At-least-once delivery and idempotency",
        "text": (
            "Event Grid promises at-least-once delivery - that's the runtime reality. What scenarios "
            "cause Event Grid to redeliver the same event to a subscriber? How does StateWriter handle "
            "a duplicate VehicleUpdate event today (from SP1-06)? What concrete bug would manifest if "
            "StateWriter just blindly upserted every event it received?"
        ),
        "answers": [
            "Important framing: at-least-once is a property of how Event Grid delivers TO subscribers, not how publishers behave. (Developer initially conflated publisher-side duplicates with delivery semantics - those are separate concerns.)",
            "Cause of redelivery #1: subscriber returns 5xx (or any non-2xx). Event Grid retries with exponential backoff for up to 24h. Common when StateWriter throws due to Cosmos throttling - the first attempt may have written before throwing, producing a duplicate on retry.",
            "Cause of redelivery #2: subscriber timeout. If StateWriter takes longer than Event Grid's delivery timeout (default 30s), Event Grid assumes failure and retries; StateWriter may have completed in the background -> duplicate.",
            "Cause of redelivery #3: acknowledgment lost in transit. StateWriter completes successfully and returns 200, but the response is lost between StateWriter and Event Grid -> Event Grid retries -> duplicate write attempt.",
            "Cause of redelivery #4: Event Grid's own internal replication. Very rare, but the service replicates events internally for durability and that can occasionally result in fanout duplicates.",
            "Rule of thumb: any subscriber that occasionally returns slow or fails will see redelivery. Not 'if' but 'when' at scale.",
            "Publisher-side duplicates (Poller publishing the same event twice due to a TfNSW glitch) are a separate concern, but the same StateWriter idempotency mechanism handles both scenarios identically.",
            "StateWriter handling (the 'stale-write guard' pattern, from SP1-06): read existing VehicleDocument by id (vehicleId) and partition key (routeShortName). If the stored document exists AND stored.Timestamp >= incoming.Timestamp, drop the event, return null (no upsert, no SignalR broadcast). Else upsert and broadcast.",
            "Cost of the guard: 1 extra Cosmos RU per invocation (the pre-read). Per SP1-06 notes: acceptable to prevent out-of-order writes from Event Grid's at-least-once delivery guarantee.",
            "Why >= and not == : the single check handles BOTH failure modes - pure duplicate (same timestamp -> skip; idempotent) AND out-of-order delivery (older event arrives after newer one -> >= triggers -> skip; prevents stale overwrite).",
            "Compare to AlerterFunction (SP1-07): no stale-write guard needed there. Alert upserts are idempotent by alertId (re-delivery overwrites with identical data; no out-of-order concern since alerts don't update in fast succession).",
            "Bug without idempotency #1 - stale data on the map: out-of-order delivery overwrites newer with older. Vehicle shows on the wrong block. Worst user-visible failure.",
            "Bug without idempotency #2 - Cosmos churn: extra writes burn RUs. cost-model.md: 'Hot loops without throttling can spike a $20+ bill in hours.' A duplicate storm during an Event Grid retry burst could push you over budget fast.",
            "Bug without idempotency #3 - SignalR broadcast storm: every duplicate event triggers a broadcast to the vehicles group. Connected dashboards see the same vehicle update flash multiple times on the map - flickering UI. SignalR Free SKU has a 20k messages/day cap (ADR-0008); a duplicate storm could exhaust the cap mid-day, killing live updates for the rest of the day.",
            "Bug without idempotency #4 - 429 throttling cascade: high duplicate rate -> Cosmos returns 429 (RU limit exceeded) -> StateWriter throws -> Event Grid retries -> more duplicates -> more 429s -> death spiral. The guard is what stops this from becoming an outage.",
            "Net: idempotency isn't 'nice to have' - it's the line of defence between 'Event Grid retried once' and 'the whole pipeline melts down under load'.",
        ],
    },
]


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"Doc not found: {DOC_PATH}")

    doc = Document(str(DOC_PATH))

    # Update the PollerFunction section title (drop the "pending" suffix).
    renamed = False
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text == TITLE_OLD:
            # Replace text while preserving the Heading 1 style.
            # Clearing runs then setting the first run's text is the safe way.
            for run in list(p.runs):
                run.text = ""
            if p.runs:
                p.runs[0].text = TITLE_NEW
            else:
                p.add_run(TITLE_NEW)
            renamed = True
            break
    if not renamed:
        print(f"WARNING: title '{TITLE_OLD}' not found; section header unchanged.")

    # Append each question block at the end of the doc.
    for q in QUESTIONS:
        doc.add_paragraph(q["heading"], style="Heading 2")
        doc.add_paragraph(q["text"], style="Normal")
        doc.add_paragraph("Model Answer", style="Heading 3")
        for bullet in q["answers"]:
            doc.add_paragraph(bullet, style="List Bullet")
        doc.add_paragraph("", style="Normal")  # spacer between questions

    doc.save(str(DOC_PATH))
    print(f"Appended {len(QUESTIONS)} question blocks to {DOC_PATH.name}")
    print(f"  Sections: Q5(a), Q5(a) follow-up, Q5(b), Q5(c), Q5(d), Q6(a), Q6(b), Q6(c)")


if __name__ == "__main__":
    main()
