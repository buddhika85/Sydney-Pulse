# Fixes the 6 daily-routine steps in the study guide that fell back to Normal style
# because "List Number" wasn't defined in the doc. Converts them to List Bullet with
# numeric prefixes so they render as a clean ordered list.
# Already executed on 2026-06-03; retained as audit trail.

from docx import Document
from pathlib import Path

DOC_PATH = Path(r"C:\BUDDHIKA\SydPulse-P6\SP1-14-Quiz-VehicleUpdate-ServiceAlert.docx")

# Exact texts of the 6 daily steps as they currently exist in the doc (no numbers yet).
# Order matters - they will be numbered 1 through 6 in this order.
DAILY_STEPS = [
    "Open the doc at a random question.",
    "Cover the Model Answer with your hand or scroll past it.",
    "Read just the question text.",
    "Speak your answer out loud as if you're in an interview.",
    "Uncover the model answer. Mentally note one thing you missed.",
    "Stop. Don't re-read. Move on.",
]


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"Doc not found: {DOC_PATH}")

    doc = Document(str(DOC_PATH))

    # Match each step by exact text. Update style + prepend number in one pass.
    fixed_count = 0
    for i, original_text in enumerate(DAILY_STEPS, start=1):
        for p in doc.paragraphs:
            if p.text == original_text:
                # Clear runs and replace with numbered text - preserves paragraph identity.
                new_text = f"{i}. {original_text}"
                for run in list(p.runs):
                    run.text = ""
                if p.runs:
                    p.runs[0].text = new_text
                else:
                    p.add_run(new_text)
                # Apply List Bullet style for visual indent (matches the rest of the doc).
                p.style = doc.styles["List Bullet"]
                fixed_count += 1
                break
        else:
            print(f"WARNING: step text not found, skipped: {original_text}")

    doc.save(str(DOC_PATH))
    print(f"Fixed {fixed_count} of {len(DAILY_STEPS)} daily-routine steps in {DOC_PATH.name}")


if __name__ == "__main__":
    main()
