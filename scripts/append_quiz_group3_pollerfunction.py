"""Append SP1-14 Group 3 PollerFunction quiz section (Q1-Q4) to the Word doc.

Q1-Q4 were quizzed live and answered by the developer with Claude feedback.
Q5-Q6 were drafted but NOT asked live - they are deferred to the next session
so the developer can answer them in real time. Q5-Q6 will be appended as a
supplementary section after that session.
"""
from docx import Document

DOC_PATH = r"C:\BUDDHIKA\SydPulse-P6\SP1-14-Quiz-VehicleUpdate-ServiceAlert.docx"

SECTION_TITLE = "SP1-14: PollerFunction (Q1-Q4 of 6 - Q5-Q6 pending next session)"

QA = [
    {
        "q": (
            "PollerFunction is decorated with [TimerTrigger(\"*/30 * * * * *\")]. What does the "
            "6-field expression mean and why is there an extra field compared to UNIX CRON? Why "
            "a timer trigger at all (vs BackgroundService or external scheduler hitting an HTTP "
            "endpoint)? If Azure scales to 3 Function App worker instances, how many times does "
            "the timer fire every 30 seconds and what protects you from triple-polling TfNSW?"
        ),
        "a": [
            "6-field NCRONTAB format (Azure Functions variant): {second} {minute} {hour} {day} {month} {day-of-week}. UNIX CRON has 5 fields; NCRONTAB adds 'seconds' at the front because UNIX CRON's minimum granularity is 1 minute.",
            "*/30 * * * * * means 'every value of seconds divisible by 30' - fires at :00 and :30 of each wall-clock minute. NOT 30 seconds after deployment. If deployed at 14:23:17 the first tick is at 14:23:30.",
            "Why timer trigger: Consumption plan has no always-on process, so BackgroundService / IHostedService can't run there (would force Premium tier and defeat the cost model).",
            "Why timer trigger: built-in distributed singleton (see below) - exactly-once-per-tick scheduling without writing leader election yourself.",
            "Why timer trigger: no external dependency. HTTP-trigger + external CRON service adds an extra moving part (scheduler, retries, monitoring). Timer keeps it inside the Functions runtime.",
            "Scaling answer: exactly 1 instance fires per tick, not 3. Mechanism is the blob lease singleton: Functions runtime acquires a lease on a blob in AzureWebJobsStorage. Only the worker holding the lease fires the timer. Others see 'lease held' and skip.",
            "Mental model: SemaphoreSlim is in-process coordination - useless across instances. Blob lease is cross-process / cross-instance coordination because it uses a shared external resource. Whenever you need 'exactly one of N processes does this', you need a shared external lock (Storage blob lease, Redis lock, Cosmos lease container, DB row lock, etc.).",
            "Separate concern: timer overlap. If a tick takes 45s but fires every 30s, the next tick starts before the previous finishes. The blob lease coordinates WHO fires, not WHETHER the previous run is done. Design for idempotent execution or pick a cadence longer than worst-case runtime.",
        ],
    },
    {
        "q": (
            "When PollerFunction has the vehicle positions and alerts decoded in memory, it "
            "publishes CloudEvents to Event Grid rather than writing to Cosmos / SignalR directly. "
            "Why? What does this fan-out buy you architecturally? What are the failure semantics "
            "if Cosmos is briefly unavailable in the Event Grid path vs. a direct-write path? "
            "What overhead are you paying for the decoupling, and when wouldn't you reach for "
            "this pattern?"
        ),
        "a": [
            "Single Responsibility Principle: PollerFunction polls + publishes; persistence (Cosmos), broadcast (SignalR), and archival (Data Lake) are downstream concerns owned by separate functions.",
            "Fan-out: one event triggers multiple consumers. VehicleUpdate.v1 -> StateWriter (Cosmos + SignalR) + Archiver. ServiceAlert.v1 -> Service Bus topic -> Alerter (Cosmos + SignalR) + Archiver.",
            "Extensibility: adding an SMS-on-alert consumer doesn't touch PollerFunction. Subscribe a new function to the existing event - open/closed principle.",
            "Independent testability: each consumer unit-tested against mocked event input. PollerFunction tests don't need Cosmos or SignalR.",
            "Independent scaling: PollerFunction (poll throughput) scales separately from StateWriter (Cosmos write throughput).",
            "Failure decoupling - the biggest win. Direct-write path during Cosmos blip: Cosmos SDK throws, poll fails mid-flight, in-memory vehicle list is lost (GC'd when function returns), timer doesn't retry the tick, 30s window of data lost.",
            "Event Grid path during Cosmos blip: poll returns 200 OK from Event Grid, PollerFunction exits clean. Event Grid attempts webhook to StateWriter, fails, retries with exponential backoff for up to 24h before dead-lettering. Data survives the Cosmos outage.",
            "Service Bus path for alerts adds more resilience: SB persists messages in durable storage. Alerter can be down for hours; messages queue then drain on recovery. MaxDeliveryCount + DLQ auto-handle poison messages.",
            "Important name correction: this is Event Grid (event routing) + Service Bus (queued delivery). NOT Event Hub - that's a different service for high-throughput stream ingestion.",
            "At-least-once delivery is now a contract. Consumers MUST be idempotent (see Q6 for StateWriter's stale-write guard).",
            "Costs - latency: end-to-end ~200-700ms vs ~50-100ms direct. Acceptable for live dashboard, not for sub-50ms requirements.",
            "Costs - money: Event Grid ~$0.60 per million events; Service Bus operations sub-cent per 1000. Trivial at our scale.",
            "Costs - operational complexity: more services to monitor; risks like exact-match drift between code and Bicep filters (see Q3 silent-failure mode).",
            "When NOT to use this pattern: single consumer (no fan-out benefit), strict sub-50ms latency, system small enough that operational complexity outweighs decoupling.",
        ],
    },
    {
        "q": (
            "Events carry type strings like com.sydneypulse.VehicleUpdate.v1. SP1-05 notes that "
            "these must match the includedEventTypes filter values in messaging.bicep exactly - "
            "any drift silently drops events. Why this format (reverse-DNS + PascalCase + .v1)? "
            "What does 'silently dropped' look like operationally if you mismatch them? Why "
            "version at v1, and how do v1 and v2 coexist during a schema migration?"
        ),
        "a": [
            "Format: com.sydneypulse.VehicleUpdate.v1 - reverse-DNS prefix + PascalCase event name + version suffix.",
            "Reverse-DNS prefix: globally unique namespace, prevents collisions across organizations. Same pattern as Java packages, MIME types, etc.",
            "PascalCase event name: semantic identifier; convention in the .NET ecosystem.",
            "Version suffix: enables schema evolution by carrying explicit version in the event identifier.",
            "Convention source: this is the CloudEvents v1.0 spec (CNCF standard at cloudevents.io), not Event Grid-specific. Any system that adopts CloudEvents (AWS EventBridge, Knative, Kafka CloudEvents serdes) uses the same shape. Portability is the real value.",
            "Source vs Type are separate CloudEvents fields. Source (/sydney-pulse/poller) identifies WHO emitted. Type identifies WHAT kind. Both filterable in subscriptions.",
            "Silent-failure mode: if you rename C# to VehicleUpdated.v1 but Bicep still filters on VehicleUpdate.v1, PublishCloudEventsAsync returns 200 OK. Event Grid accepts the event into the topic, evaluates subscription filters, finds no match, discards the event silently.",
            "What you observe operationally: PollerFunction logs HEALTHY (publish succeeded). StateWriter invocations drop to zero. Cosmos writes stop. SignalR broadcasts stop. Dashboard freezes silently. No errors anywhere - every health check stays green.",
            "Failure manifests as absence of activity, not as error. This is the brutal pattern of Azure event-driven systems - trust delta-against-baseline metrics more than absence-of-errors.",
            "How to catch in production: heartbeat metric (StateWriter invocations/min < N triggers alert); Event Grid metric PublishedEvents vs DeliveredEvents divergence; Cosmos write-volume metric; synthetic check on /api/vehicles asserting receivedAt is recent.",
            "Versioning rules - backwards-compatible (no version bump): adding an optional field. Breaking (needs v2): adding a required field, removing a field, renaming a field, changing semantics.",
            "Migration / coexistence: dual-write phase. PollerFunction publishes BOTH v1 and v2 versions of each event for a transition period. New consumers subscribe to v2; old consumers stay on v1. After all consumers migrated, stop publishing v1, remove v1 subscriptions and Bicep filters.",
            "Alternative coexistence model: an adapter function subscribes to v2 and re-publishes v1 for legacy consumers. Less duplication but more moving parts.",
            "Lifecycle: producer dual-writes -> migrate consumers one at a time -> producer drops v1 last. Roll-out is producer-first, consumer-migration-second, producer-cleanup-last.",
        ],
    },
    {
        "q": (
            "Suppose the polling loop iterates [buses, trains, ferries, lightrail, metro] and on "
            "one tick the ferries call throws (TfNSW returned a 500 that exhausted Polly's "
            "retry budget). What does the current loop do? Does the timer runtime retry the "
            "tick? What happens to modes that were already published, and what happens to modes "
            "after ferries? How would you make this resilient, and what's the trade-off between "
            "fail-fast and continue-on-error?"
        ),
        "a": [
            "Current behavior: the loop has only one try boundary - the function itself. When await PublishVehicleUpdatesAsync('ferries') throws, the exception propagates out of the foreach. The loop is interrupted. metro is NEVER invoked on this tick.",
            "Polly is NOT a loop-level retry mechanism. Polly operates at the HttpClient layer, scoped to a SINGLE HTTP request. When EnsureSuccessStatusCode throws, Polly has ALREADY given up - either retries exhausted or status was non-transient (401, 403, 404, etc.). Loop-level retry must be written explicitly if wanted.",
            "Timer trigger does NOT retry the tick. This differs from other triggers: Service Bus redelivers up to MaxDeliveryCount then DLQ; Queue Storage similar; Event Grid (consumer side) retries with exponential backoff up to 24h. Timer just waits for the next scheduled tick.",
            "No distributed transaction / rollback. Events published before the throw stay on Event Grid - they're not undone. PollerFunction can produce a partially-completed tick: some modes' events on Event Grid, others not.",
            "Specific outcome when ferries fails: modes before ferries (buses, trains, lightrail) - events on Event Grid, normal downstream flow. ferries - failed, no events. metro - never invoked. 30s of metro updates lost for this tick.",
            "Next tick (30s later) fetches fresh current state. Data loss is bounded for current-state systems like this - we lose 30s of metro updates, not unrecoverable history.",
            "If the payload were event-sourced deltas (e.g., 'passenger boarded/alighted') the loss would be permanent. Our system is current-state, so the failure mode is bounded.",
            "Functions runtime observability when the function throws: web UI marks the invocation as Failed; App Insights captures the exception with stack trace, request telemetry, and dependency telemetry.",
            "Resilient redesign: per-mode try/catch inside the foreach.",
            "Resilient redesign: structured logging with mode field - logger.LogError(ex, 'Mode {Mode} poll failed', mode).",
            "Resilient redesign: custom metric like tfnsw_poll_failures{mode=ferries} so failures surface on dashboards. Without metrics, swallowed exceptions become invisible - the same silent-failure trap as Q3.",
            "Resilient redesign: continue the loop after catching, so subsequent modes still publish.",
            "Resilient redesign optional: aggregate caught exceptions and re-throw at end of RunAsync if you want the overall tick marked failed in App Insights. Otherwise the function reports success even when N modes failed.",
            "Trade-off fail-fast (current): first failure surfaces immediately as a failed invocation. Hard to ignore. Cost: loses subsequent modes for that tick.",
            "Trade-off continue-on-error (proposed): max data capture; partial tick still useful. Cost: failures can be silently swallowed without good logging and metrics.",
            "Decision rule: for poll-current-state systems, continue-on-error usually wins (failure mode is bounded staleness, not permanent loss). For event-sourced delta systems, stricter approach (per-event ack, queue with redelivery) is mandatory.",
        ],
    },
]


def main() -> None:
    doc = Document(DOC_PATH)

    # Drop trailing blank paragraphs so the new section starts cleanly.
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)

    doc.add_paragraph("")  # separator
    doc.add_paragraph(SECTION_TITLE, style="Heading 1")

    for i, item in enumerate(QA, start=1):
        doc.add_paragraph(f"Question {i}", style="Heading 2")
        doc.add_paragraph(item["q"], style="Normal")
        doc.add_paragraph("Model Answer", style="Heading 3")
        for bullet in item["a"]:
            doc.add_paragraph(bullet, style="List Bullet")
        doc.add_paragraph("")  # separator after each Q

    doc.save(DOC_PATH)
    print(f"Saved {DOC_PATH}")
    print(f"Added section: {SECTION_TITLE} with {len(QA)} questions.")


if __name__ == "__main__":
    main()
